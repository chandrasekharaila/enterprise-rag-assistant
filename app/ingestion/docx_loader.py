from docx import Document as DocxDocument

from app.ingestion.base_loader import BaseLoader
from app.schemas.document import Document


class DOCXLoader(BaseLoader):

    def load(self) -> Document:

        doc = DocxDocument(self.file_path)

        paragraphs = []

        for paragraph in doc.paragraphs:
            paragraphs.append(paragraph.text)

        content = "\n".join(paragraphs)

        metadata = {
            "source": self.file_path.name,
            "type": "docx",
            "paragraphs": len(doc.paragraphs),
        }

        return Document(
            content=content,
            metadata=metadata,
        )